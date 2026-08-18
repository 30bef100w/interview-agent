"""面试报告导出为 Word（.docx）文档。"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

INDIGO = (79, 70, 229)
GREEN = (5, 150, 105)
RED = (220, 38, 38)
AMBER = (217, 119, 6)
DARK = (39, 39, 42)
GRAY = (113, 113, 122)


def _run(p, text: str, size: float = 10.5, bold: bool = False, color=DARK) -> None:
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if color:
        run.font.color.rgb = RGBColor(*color)


def _para(doc: Document, text: str = "", size: float = 10.5, bold: bool = False,
          color=DARK, align=None, space_after: float = 6) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        _run(p, text, size, bold, color)
    return p


def _score_color(v) -> tuple[int, int, int]:
    if v >= 8:
        return GREEN
    if v >= 6:
        return AMBER
    return RED


def _add_list(doc: Document, items: list, bullet: bool = True) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
        _run(p, str(item))


def build_report_docx(report: dict, meta: dict) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Pt(56)
        section.bottom_margin = Pt(56)
        section.left_margin = Pt(64)
        section.right_margin = Pt(64)

    _para(doc, "AI 模拟面试报告", size=22, bold=True, color=INDIGO,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(doc, f"面试形态：{meta.get('mode_label', '')} · 生成时间：{meta.get('created_at', '')}",
          size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    dims = report.get("dimension_scores") or {}
    if dims:
        total = sum(dims.values()) / len(dims)
        _para(doc, f"综合评分：{total:.1f} / 10", size=16, bold=True,
              color=_score_color(total), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        cells = table.rows[0].cells
        _run(cells[0].paragraphs[0], "维度", bold=True, color=DARK)
        _run(cells[1].paragraphs[0], "得分", bold=True, color=DARK)
        for k, v in dims.items():
            row = table.add_row().cells
            _run(row[0].paragraphs[0], str(k))
            _run(row[1].paragraphs[0], f"{v} / 10", bold=True, color=_score_color(v))
        doc.add_paragraph()

    _para(doc, "总体评价", size=14, bold=True, color=INDIGO, space_after=4)
    _para(doc, str(report.get("summary", "")), size=10.5, space_after=12)

    per_question = report.get("per_question") or []
    if per_question:
        _para(doc, "逐题作答详情", size=14, bold=True, color=INDIGO, space_after=6)
        for i, q in enumerate(per_question):
            _para(doc, f"第 {i + 1} 题 · {q.get('topic', '')}（{q.get('score', '-')}/10）",
                  size=12, bold=True, color=_score_color(float(q.get("score") or 0)), space_after=3)
            _para(doc, f"问题：{q.get('question', '')}", size=10.5, space_after=3)
            answers = q.get("my_answers") or []
            _para(doc, "我的作答：", size=10.5, bold=True, space_after=2)
            for a in answers:
                _para(doc, f"· {a}", size=10.5, color=GRAY, space_after=2)
            feedback = q.get("feedback")
            if feedback:
                _para(doc, f"AI 点评：{feedback}", size=10.5, space_after=2)
            reference = q.get("reference_answer")
            if reference:
                _para(doc, f"参考答案：{reference}", size=10.5, color=INDIGO, space_after=2)
            doc.add_paragraph()

    strengths = report.get("strengths") or []
    weaknesses = report.get("weaknesses") or []
    if strengths or weaknesses:
        if strengths:
            _para(doc, "优点", size=14, bold=True, color=GREEN, space_after=4)
            _add_list(doc, strengths)
        if weaknesses:
            _para(doc, "待提升", size=14, bold=True, color=RED, space_after=4)
            _add_list(doc, weaknesses)

    suggestions = report.get("suggestions") or []
    if suggestions:
        _para(doc, "提升建议", size=14, bold=True, color=INDIGO, space_after=4)
        _add_list(doc, suggestions, bullet=False)

    from io import BytesIO

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
