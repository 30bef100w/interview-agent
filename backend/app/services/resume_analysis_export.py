"""简历分析报告导出 Word / PDF。"""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from fpdf import FPDF

SKY = (2, 132, 199)
DARK = (39, 39, 42)
GRAY = (113, 113, 122)

FONT_CANDIDATES = [
    Path(r"C:\Windows\Fonts\msyh.ttc"),
    Path(r"C:\Windows\Fonts\msyh.ttf"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),
    Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"),
    Path("/System/Library/Fonts/PingFang.ttc"),
]


def _run(p, text: str, size: float = 10.5, bold: bool = False, color=DARK) -> None:
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if color:
        run.font.color.rgb = RGBColor(*color)


def _para(doc: Document, text: str = "", size: float = 10.5, bold: bool = False,
          color=DARK, align=None, space_after: float = 6) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        _run(p, text, size, bold, color)


def _add_list(doc: Document, items: list) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        _run(p, str(item))


def build_analysis_docx(analysis: dict, meta: dict) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Pt(56)
        section.bottom_margin = Pt(56)
        section.left_margin = Pt(64)
        section.right_margin = Pt(64)

    _para(doc, "AI 简历分析报告", size=20, bold=True, color=SKY,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(
        doc,
        f"简历：{meta.get('filename', '')} · 生成时间：{meta.get('created_at', '')}",
        size=9,
        color=GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )

    score = analysis.get("score") or 0
    _para(doc, f"竞争力评分：{score} / 10", size=14, bold=True, color=SKY, space_after=10)

    _para(doc, "总体评价", size=13, bold=True, color=SKY, space_after=4)
    _para(doc, str(analysis.get("summary") or "（无）"), size=10.5, space_after=10)

    for title, key in (
        ("优势", "strengths"),
        ("风险 / 短板", "risks"),
        ("优化建议", "improvements"),
        ("面试可能深挖", "interview_focus"),
    ):
        items = analysis.get(key) or []
        if not items:
            continue
        _para(doc, title, size=13, bold=True, color=SKY, space_after=4)
        _add_list(doc, items)
        doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class _AnalysisPdf(FPDF):
    def footer(self) -> None:
        self.set_y(-12)
        self.set_font("Body", size=8)
        self.set_text_color(120, 120, 120)
        self.cell(0, 8, f"{self.page_no()}", align="C")


def build_analysis_pdf(analysis: dict, meta: dict) -> bytes:
    font_path = next((p for p in FONT_CANDIDATES if p.exists()), None)
    pdf = _AnalysisPdf(format="A4", unit="mm")
    pdf.set_auto_page_break(auto=True, margin=16)
    pdf.add_page()

    if font_path is None:
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 8, "Chinese font not found.")
        return bytes(pdf.output())

    pdf.add_font("Body", style="", fname=str(font_path))
    pdf.add_font("Body", style="B", fname=str(font_path))

    pdf.set_font("Body", style="B", size=18)
    pdf.set_text_color(*SKY)
    pdf.cell(0, 10, "AI 简历分析报告", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Body", size=9)
    pdf.set_text_color(*GRAY)
    pdf.cell(
        0,
        6,
        f"简历：{meta.get('filename', '')} · 生成时间：{meta.get('created_at', '')}",
        align="C",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.ln(4)

    pdf.set_font("Body", style="B", size=13)
    pdf.set_text_color(*SKY)
    pdf.cell(0, 8, f"竞争力评分：{analysis.get('score') or 0} / 10", new_x="LMARGIN", new_y="NEXT")

    def section(title: str) -> None:
        pdf.set_font("Body", style="B", size=12)
        pdf.set_text_color(*SKY)
        pdf.cell(0, 8, title, new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(*DARK)
        pdf.set_font("Body", size=10)

    section("总体评价")
    pdf.multi_cell(0, 6, str(analysis.get("summary") or "（无）"))
    pdf.ln(2)

    for title, key in (
        ("优势", "strengths"),
        ("风险 / 短板", "risks"),
        ("优化建议", "improvements"),
        ("面试可能深挖", "interview_focus"),
    ):
        items = analysis.get(key) or []
        if not items:
            continue
        section(title)
        for item in items:
            pdf.multi_cell(0, 5.5, f"· {item}")
        pdf.ln(1)

    return bytes(pdf.output())
